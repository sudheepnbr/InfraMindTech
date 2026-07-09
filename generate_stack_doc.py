"""Generate InfraMindTech website structure & stack Word document."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h1(text):
    doc.add_heading(text, level=1)


def add_h2(text):
    doc.add_heading(text, level=2)


def add_p(text):
    doc.add_paragraph(text)


def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph("")


add_title("InfraMindTech Website")
add_p("Structure & Technology Stack Documentation")
add_p("Generated: July 10, 2026")
doc.add_paragraph("")

add_h1("1. Live URLs")
add_table(
    ["Role", "URL"],
    [
        ["Public website", "https://sudheepnbr.github.io/InfraMindTech/"],
        ["Backend + CMS API", "https://inframindtech.onrender.com/"],
        ["Admin panel", "https://inframindtech.onrender.com/admin/"],
        ["GitHub repository", "https://github.com/sudheepnbr/InfraMindTech"],
    ],
)

add_h1("2. Architecture Overview")
add_p("The project uses a split deployment model:")
add_bullets(
    [
        "GitHub (main branch) stores the full codebase.",
        "GitHub Actions deploys frontend-only files to the gh-pages branch for public hosting.",
        "Render auto-deploys from main and runs the Python Flask backend, CMS admin, and content API.",
        "GitHub Pages serves the static public website. Render serves backend, admin, and API.",
    ]
)

add_h1("3. Frontend Stack")
add_table(
    ["Layer", "Technology"],
    [
        ["Markup", "HTML5 (static pages)"],
        ["CSS", "Custom CSS + Bootstrap 5.3.3"],
        ["JavaScript", "Vanilla JS (no React/Vue/Angular)"],
        ["Icons", "Font Awesome 6.5.1"],
        ["Fonts", "Google Fonts (Inter)"],
        ["Animations", "AOS 2.3.4 + GSAP 3.12.5 (home hero)"],
        ["Layout", "Shared partials loaded via includes.js"],
    ],
)

add_h2("3.1 CDN Dependencies")
add_bullets(["Bootstrap 5.3.3", "Font Awesome 6.5.1", "AOS 2.3.4", "GSAP 3.12.5 (home page only)"])

add_h2("3.2 Key Frontend Files")
add_table(
    ["File / Folder", "Purpose"],
    [
        ["index.html", "Home page"],
        ["solutions/, services/, products/, industries/, resources/, about/, contact/", "Inner pages with breadcrumbs"],
        ["partials/header.html, partials/footer.html", "Shared site chrome"],
        ["js/site-base.js", "Sets base URL for GitHub Pages paths"],
        ["js/includes.js", "Loads header/footer, navigation links, active nav state"],
        ["js/content-loader.js", "Loads CMS content from API or content.json"],
        ["js/app.js", "Theme toggle, mobile nav, FAQ, counters, page transitions"],
        ["data/content.json", "Editable site text (fallback + CMS source)"],
        ["css/style.css, css/responsive.css", "Main stylesheets"],
    ],
)

add_h1("4. Backend Stack (Render Only)")
add_table(
    ["Layer", "Technology"],
    [
        ["Language", "Python 3.11"],
        ["Framework", "Flask 3.x"],
        ["Server", "Gunicorn"],
        ["Config", "render.yaml"],
        ["Entry point", "server.py"],
        ["Content storage", "data/content.json (no database)"],
    ],
)

add_h2("4.1 Python Dependencies")
add_bullets(["flask>=3.0.0", "gunicorn>=21.0.0"])

add_h2("4.2 Backend Features")
add_bullets(
    [
        "REST API for CMS content (GET/PUT /api/content)",
        "Admin authentication (POST /api/login)",
        "Media uploads (POST /api/upload)",
        "Session-based admin auth",
        "CORS enabled for GitHub Pages",
        "Serves static site and admin panel on Render",
    ]
)

add_h1("5. CMS / Admin Panel")
add_table(
    ["Item", "Detail"],
    [
        ["Admin UI", "admin/index.html, admin.js, admin.css"],
        ["Content store", "data/content.json"],
        ["Local development", "python server.py → http://localhost:8080"],
        ["Local admin URL", "http://localhost:8080/admin/"],
        ["Environment config", ".env (not committed to git)"],
    ],
)

add_h1("6. Deployment")
add_h2("6.1 GitHub Pages (Public Website)")
add_bullets(
    [
        "Workflow: .github/workflows/deploy-pages.yml",
        "Trigger: Push to main branch",
        "Output: Static frontend files published to gh-pages branch",
        "GitHub Pages setting: Branch gh-pages, folder / (root)",
        "Excluded from public deploy: server.py, admin/, requirements.txt, render.yaml",
    ]
)

add_h2("6.2 Render (Backend + CMS)")
add_bullets(
    [
        "Build command: pip install -r requirements.txt",
        "Start command: gunicorn server:app --bind 0.0.0.0:$PORT",
        "Plan: Free tier",
        "Environment variables: ADMIN_USERNAME, ADMIN_PASSWORD, SECRET_KEY",
    ]
)

add_h1("7. Website Page Structure")
add_table(
    ["Page", "URL Path", "Breadcrumb"],
    [
        ["Home", "/", "—"],
        ["Solutions", "/solutions/", "Home > Solutions"],
        ["Services", "/services/", "Home > Services"],
        ["Products", "/products/", "Home > Products"],
        ["Industries", "/industries/", "Home > Industries"],
        ["Resources", "/resources/", "Home > Resources"],
        ["About", "/about/", "Home > About"],
        ["Contact", "/contact/", "Home > Contact"],
    ],
)
add_p("Old .html URLs (example: products.html) redirect to clean folder URLs.")

add_h1("8. Project Folder Structure")
add_p(
    "E:\\Website\\\n"
    "├── index.html                     Home page\n"
    "├── solutions/ services/ products/ Inner pages\n"
    "├── industries/ resources/ about/ contact/\n"
    "├── partials/                      Shared header & footer\n"
    "├── css/                           Stylesheets\n"
    "├── js/                            JavaScript modules\n"
    "├── images/                        Logo and media assets\n"
    "├── data/                          content.json\n"
    "├── admin/                         CMS dashboard (Render)\n"
    "├── server.py                      Flask backend (Render)\n"
    "├── requirements.txt               Python dependencies\n"
    "├── render.yaml                    Render deployment config\n"
    "├── .github/workflows/             GitHub Pages deploy action\n"
    "└── backend-local/                 Local backup (gitignored)"
)

add_h1("9. Navigation & Branding")
add_bullets(
    [
        "Header navigation links to dedicated pages with breadcrumbs.",
        "Logo shows iMT icon + InfraMindTech name + tagline: INNOVATE • INTEGRATE • INSPIRE",
        "Inner pages use page-header, AOS animations, and page-enter transition.",
        "Content on GitHub Pages loads from Render API when available.",
    ]
)

add_h1("10. Stack Summary")
add_table(
    ["Category", "Choice"],
    [
        ["Project type", "Static frontend + JSON-backed CMS backend"],
        ["Frontend", "HTML / CSS / Vanilla JavaScript"],
        ["Backend", "Python Flask"],
        ["Hosting", "GitHub Pages + Render"],
        ["Database", "None (JSON file storage)"],
        ["Build tools", "None (no Webpack/Vite/npm)"],
        ["Version control", "Git + GitHub"],
    ],
)

out_docx = Path(__file__).parent / "InfraMindTech-Website-Structure-and-Stack.docx"
doc.save(out_docx)
print(f"Created: {out_docx}")

try:
    import win32com.client  # type: ignore

    out_doc = Path(__file__).parent / "InfraMindTech-Website-Structure-and-Stack.doc"
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    docx_path = str(out_docx.resolve())
    doc_path = str(out_doc.resolve())
    wdoc = word.Documents.Open(docx_path)
    wdoc.SaveAs2(doc_path, FileFormat=0)
    wdoc.Close()
    word.Quit()
    print(f"Created: {out_doc}")
except Exception as exc:
    print(f"Note: .doc conversion skipped ({exc}). Use the .docx file in Microsoft Word.")
